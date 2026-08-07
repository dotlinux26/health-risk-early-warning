"""Xuất file mẫu TỪ DỮ LIỆU THẬT NHANES (CDC) — đúng định dạng hệ thống đọc được.

Chạy:
    python3 scripts/export_nhanes_samples.py

Với mỗi ca thật được chọn (khỏe mạnh / tăng huyết áp / đái tháo đường / suy
thận), xuất 3 loại:
    data/sample_nhanes/{pid}.docx   báo cáo khám (upload qua chat)
    data/sample_nhanes/{pid}.pdf    báo cáo khám
    data/sample_nhanes/{pid}.csv    nhật ký dạng long (schema chuẩn)

Lưu ý: NHANES là khảo sát 1 kỳ khám/người nên mỗi file = 1 ngày đo. Để tích lũy
7 ngày trong chat, nhập thêm các ngày khác bằng tay hoặc dùng sample_long.csv.
"""
from __future__ import annotations

import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import pandas as pd

OUT = Path("data/sample_nhanes")


def _pick(df: pd.DataFrame, cond: pd.Series, n: int = 1) -> pd.DataFrame:
    return df[cond].dropna(subset=["systolic_bp", "diastolic_bp"]).head(n)


def _items(row: pd.Series) -> list[tuple[str, str, str]]:
    it: list[tuple[str, str, str]] = []
    if pd.notna(row["systolic_bp"]):
        it.append(("Huyết áp tâm thu", f"{row['systolic_bp']:.0f}", "mmHg"))
    if pd.notna(row["diastolic_bp"]):
        it.append(("Huyết áp tâm trương", f"{row['diastolic_bp']:.0f}", "mmHg"))
    if pd.notna(row["heart_rate"]):
        it.append(("Nhịp tim", f"{row['heart_rate']:.0f}", "bpm"))
    if pd.notna(row["glucose_fasting"]):
        it.append(("Đường huyết lúc đói", f"{row['glucose_fasting']:.1f}", "mmol/L"))
    if pd.notna(row["hba1c"]):
        it.append(("HbA1c", f"{row['hba1c']:.1f}", "%"))
    if pd.notna(row["creatinine"]):
        it.append(("Creatinine", f"{row['creatinine']:.2f}", "mg/dL"))
    if pd.notna(row["bmi"]):
        it.append(("BMI", f"{row['bmi']:.1f}", "kg/m2"))
    return it


def _find_font() -> str:
    import glob

    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    hits = glob.glob("/usr/share/fonts/**/*Sans*.ttf", recursive=True)
    return hits[0] if hits else "helv"


def gen_docx(path: Path, pid: str, seqn: int, items: list[tuple[str, str, str]]) -> None:
    import docx

    d = docx.Document()
    d.add_heading(f"BÁO CÁO KẾT QUẢ KIỂM TRA SỨC KHỎE — {pid}", level=1)
    d.add_paragraph(f"Ngày khám: 15/06/2025 | Mã hồ sơ CDC: {int(seqn)}")
    d.add_heading("Kết quả", level=2)
    t = d.add_table(rows=len(items) + 1, cols=3)
    t.rows[0].cells[0].text = "Chỉ tiêu"
    t.rows[0].cells[1].text = "Kết quả"
    t.rows[0].cells[2].text = "Đơn vị"
    for i, (name, value, unit) in enumerate(items, start=1):
        t.rows[i].cells[0].text = name
        t.rows[i].cells[1].text = value
        t.rows[i].cells[2].text = unit
    d.save(str(path))


def gen_pdf(path: Path, pid: str, seqn: int, items: list[tuple[str, str, str]]) -> None:
    import fitz

    fontfile = _find_font()
    doc = fitz.open()
    page = doc.new_page()
    fontname = "helv"
    if fontfile != "helv":
        fontname = "F0"
        page.insert_font(fontname=fontname, fontfile=fontfile)
    y = 60
    page.insert_text((50, y), f"BÁO CÁO KẾT QUẢ KIỂM TRA SỨC KHỎE — {pid}", fontsize=14, fontname=fontname)
    y += 24
    page.insert_text((50, y), f"Ngày khám: 15/06/2025 | Mã hồ sơ CDC: {int(seqn)}", fontsize=11, fontname=fontname)
    y += 30
    for name, value, unit in items:
        page.insert_text((50, y), f"{name}: {value} {unit}", fontsize=11, fontname=fontname)
        y += 18
    doc.save(str(path))


def gen_csv(path: Path, pid: str, items: list[tuple[str, str, str]]) -> None:
    df = pd.DataFrame(
        [(pid, "2025-06-15", m, float(v)) for m, v in _metric_values(items)],
        columns=["patient_id", "timestamp", "metric", "value"],
    )
    df.to_csv(path, index=False)


def _to_metric(name: str) -> str | None:
    return {
        "Huyết áp tâm thu": "systolic_bp",
        "Huyết áp tâm trương": "diastolic_bp",
        "Nhịp tim": "heart_rate",
        "Đường huyết lúc đói": "glucose_fasting",
        "HbA1c": "hba1c",
        "Creatinine": "creatinine",
        "BMI": "bmi",
    }.get(name)


def _metric_values(items: list[tuple[str, str, str]]) -> list[tuple[str, float]]:
    out = []
    for name, value, unit in items:
        m = _to_metric(name)
        if m is not None:
            out.append((m, float(value)))
    return out


def main() -> None:
    csv = pd.read_csv("data/datasets/nhanes_2017_2018.csv")
    OUT.mkdir(parents=True, exist_ok=True)

    cases = [
        ("NOK0001", "Khỏe mạnh", _pick(csv, (csv["label"] == 0) & (csv["systolic_bp"] < 125)
                                        & (csv["hba1c"] < 5.5) & (csv["glucose_fasting"] < 5.6))),
        ("NHTN0001", "Tăng huyết áp", _pick(csv, (csv["label"] == 1)
                                            & (csv["systolic_bp"] >= 150))),
        ("NDM0001", "Đái tháo đường", _pick(csv, (csv["hba1c"] >= 7.0))),
        ("NCKD0001", "Suy thận nghi ngờ", _pick(csv, (csv["creatinine"] >= 1.5))),
    ]

    for pid, desc, row in cases:
        if row.empty:
            print(f"  BỎ QUA {pid} ({desc}) — không có ca phù hợp")
            continue
        r = row.iloc[0]
        seqn = int(r["seqn"])
        items = _items(r)
        print(f"== {pid} ({desc}) — hồ sơ CDC #{seqn}, {int(r['age'])} tuổi, "
              f"label_htn={int(r['label_htn'])}, label_dm={int(r['label_dm'])}")
        gen_docx(OUT / f"{pid}.docx", pid, seqn, items)
        gen_pdf(OUT / f"{pid}.pdf", pid, seqn, items)
        gen_csv(OUT / f"{pid}.csv", pid, items)
        print(f"   {len(items)} chỉ số -> {OUT}/{pid}.docx / .pdf / .csv")


if __name__ == "__main__":
    main()
