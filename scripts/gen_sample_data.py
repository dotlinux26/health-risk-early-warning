#!/usr/bin/env python3
"""Sinh dữ liệu mẫu: dataset CSV chuẩn + báo cáo y tế DOCX/PDF.

Chạy:
    python3 scripts/gen_sample_data.py
"""
from __future__ import annotations

import random
from pathlib import Path

import pandas as pd

random.seed(42)
OUT = Path("data")
REPORTS = OUT / "reports"


def make_patient_rows() -> list[tuple]:
    rows: list[tuple] = []

    # P001: khỏe mạnh, các chỉ số ổn định
    for d in range(120):
        ts = pd.Timestamp("2025-01-01") + pd.Timedelta(days=d)
        rows.append(("P001", ts, "systolic_bp", round(118 + random.uniform(-4, 4), 1)))
        rows.append(("P001", ts, "diastolic_bp", round(76 + random.uniform(-3, 3), 1)))
        rows.append(("P001", ts, "heart_rate", round(70 + random.uniform(-5, 5), 1)))
        rows.append(("P001", ts, "bmi", round(22.5 + random.uniform(-0.4, 0.4), 1)))

    # P002: huyết áp tăng dần trong 45 ngày cuối -> nguy cơ tim mạch
    bp = 122
    for d in range(120):
        ts = pd.Timestamp("2025-01-01") + pd.Timedelta(days=d)
        if d >= 75:
            bp += 0.32
        rows.append(("P002", ts, "systolic_bp", round(bp + random.uniform(-3, 3), 1)))
        rows.append(("P002", ts, "diastolic_bp", round(bp / 1.6 + random.uniform(-2, 2), 1)))
        rows.append(("P002", ts, "heart_rate", round(74 + random.uniform(-4, 4), 1)))

    # P003: tiểu đường (glucose + HbA1c cao)
    for d in range(90):
        ts = pd.Timestamp("2025-01-01") + pd.Timedelta(days=d)
        rows.append(("P003", ts, "glucose_fasting", round(8.1 + random.uniform(-0.5, 0.5), 2)))
        rows.append(("P003", ts, "hba1c", round(7.2 + random.uniform(-0.1, 0.1), 2)))
        rows.append(("P003", ts, "bmi", round(27.8 + random.uniform(-0.3, 0.3), 1)))

    # P004: chức năng thận suy giảm (creatinine tăng, eGFR giảm)
    cr = 0.9
    for d in range(80):
        ts = pd.Timestamp("2025-01-01") + pd.Timedelta(days=d)
        if d >= 50:
            cr += 0.012
        rows.append(("P004", ts, "creatinine", round(cr + random.uniform(-0.02, 0.02), 2)))
        rows.append(("P004", ts, "egfr", round(95 - (cr - 0.9) * 55 + random.uniform(-1, 1), 1)))

    # P005: ít dữ liệu (5 ngày) -> hệ thống báo INSUFFICIENT_DATA / đánh giá theo luật
    for d in range(5):
        ts = pd.Timestamp("2025-03-01") + pd.Timedelta(days=d)
        rows.append(("P005", ts, "systolic_bp", round(130 + random.uniform(-3, 3), 1)))
        rows.append(("P005", ts, "glucose_fasting", round(6.4 + random.uniform(-0.3, 0.3), 2)))

    return rows


def gen_csv() -> Path:
    df = pd.DataFrame(make_patient_rows(), columns=["patient_id", "timestamp", "metric", "value"])
    path = OUT / "sample_long.csv"
    df.to_csv(path, index=False)
    return path


def gen_docx(path: Path, patient: str, items: list[tuple[str, str, str]]) -> None:
    import docx

    d = docx.Document()
    d.add_heading(f"BÁO CÁO KẾT QUẢ KIỂM TRA SỨC KHỎE — {patient}", level=1)
    d.add_paragraph(f"Ngày khám: 15/06/2025 | Mã bệnh nhân: {patient}")
    d.add_heading("Kết quả", level=2)
    t = d.add_table(rows=len(items) + 1, cols=3)
    t.rows[0].cells[0].text = "Chỉ tiêu"
    t.rows[0].cells[1].text = "Kết quả"
    t.rows[0].cells[2].text = "Đơn vị"
    for i, (name, value, unit) in enumerate(items, start=1):
        t.rows[i].cells[0].text = name
        t.rows[i].cells[1].text = value
        t.rows[i].cells[2].text = unit
    d.save(str(path))


def _find_font() -> str:
    import glob

    candidates = [
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for c in candidates:
        if Path(c).exists():
            return c
    hits = glob.glob("/usr/share/fonts/**/*Sans*.ttf", recursive=True)
    return hits[0] if hits else "helv"


def gen_pdf(path: Path, patient: str, items: list[tuple[str, str, str]]) -> None:
    import fitz

    fontfile = _find_font()
    doc = fitz.open()
    page = doc.new_page()
    fontname = "helv"
    if fontfile != "helv":
        fontname = "F0"
        page.insert_font(fontname=fontname, fontfile=fontfile)

    y = 60
    page.insert_text((50, y), f"BÁO CÁO KẾT QUẢ KIỂM TRA SỨC KHỎE — {patient}", fontsize=14, fontname=fontname)
    y += 24
    page.insert_text((50, y), f"Ngày khám: 15/06/2025 | Mã bệnh nhân: {patient}", fontsize=11, fontname=fontname)
    y += 30
    for name, value, unit in items:
        page.insert_text((50, y), f"{name}: {value} {unit}", fontsize=11, fontname=fontname)
        y += 18
    doc.save(str(path))


def gen_reports() -> list[Path]:
    REPORTS.mkdir(parents=True, exist_ok=True)
    sets = {
        "report_huyet_ap.docx": (
            "P002",
            [("Huyết áp tâm thu", "152", "mmHg"), ("Huyết áp tâm trương", "94", "mmHg"),
             ("Nhịp tim", "88", "bpm"), ("Cân nặng", "78", "kg"), ("BMI", "26.4", "kg/m2")],
        ),
        "report_tieu_duong.docx": (
            "P003",
            [("Glucose lúc đói", "8.4", "mmol/L"), ("HbA1c", "7.3", "%"),
             ("Creatinine", "1.1", "mg/dL"), ("BMI", "28.1", "kg/m2")],
        ),
        "report_than.docx": (
            "P004",
            [("Creatinine", "1.5", "mg/dL"), ("eGFR", "52", "mL/min/1.73m2"),
             ("Huyết áp tâm thu", "141", "mmHg"), ("Huyết áp tâm trương", "89", "mmHg")],
        ),
        "report_khoe.docx": (
            "P001",
            [("Huyết áp tâm thu", "119", "mmHg"), ("Huyết áp tâm trương", "75", "mmHg"),
             ("Nhịp tim", "68", "bpm"), ("Glucose lúc đói", "5.2", "mmol/L"),
             ("HbA1c", "5.4", "%"), ("BMI", "22.6", "kg/m2")],
        ),
        "report_huyet_ap.pdf": (
            "P002",
            [("Huyết áp tâm thu", "150", "mmHg"), ("Huyết áp tâm trương", "92", "mmHg"),
             ("Nhịp tim", "86", "bpm")],
        ),
    }
    paths: list[Path] = []
    for fname, (pid, items) in sets.items():
        p = REPORTS / fname
        if fname.endswith(".pdf"):
            gen_pdf(p, pid, items)
        else:
            gen_docx(p, pid, items)
        paths.append(p)
    return paths


if __name__ == "__main__":
    OUT.mkdir(parents=True, exist_ok=True)
    csv_path = gen_csv()
    report_paths = gen_reports()
    print(f"CSV: {csv_path}")
    for p in report_paths:
        print(f"Báo cáo: {p}")
