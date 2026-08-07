"""Trích xuất văn bản từ PDF và DOCX."""
from __future__ import annotations

from pathlib import Path


def extract_text(path: str | Path) -> list[str]:
    """Trả về danh sách các dòng văn bản từ file (pdf/docx/doc)."""
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix in {".docx", ".doc"}:
        return _extract_docx(path)
    if suffix == ".txt":
        return path.read_text(encoding="utf-8", errors="ignore").splitlines()
    raise ValueError(f"Định dạng không hỗ trợ: {suffix} (hỗ trợ pdf, docx, txt)")


def _extract_pdf(path: Path) -> list[str]:
    import fitz  # PyMuPDF

    lines: list[str] = []
    with fitz.open(path) as doc:
        for page in doc:
            lines.extend(page.get_text("text").splitlines())
    return lines


def _extract_docx(path: Path) -> list[str]:
    import docx

    d = docx.Document(str(path))
    lines: list[str] = [p.text for p in d.paragraphs]
    # Bảng xét nghiệm: ghép các ô trong cùng một hàng thành một dòng để
    # parser khớp "Tên chỉ số + giá trị + đơn vị" trên cùng dòng.
    for table in d.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return [l.strip() for l in lines if l.strip()]
